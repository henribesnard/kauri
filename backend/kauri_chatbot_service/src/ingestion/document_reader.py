"""
Document Reader - Read documents from base_connaissances with table support
"""
from pathlib import Path
from typing import Dict, Any, List
import structlog
import docx
from docx.table import Table
from docx.text.paragraph import Paragraph

logger = structlog.get_logger()


class DocumentReader:
    """
    Read and parse documents from base_connaissances directory
    Preserves structure including tables (converted to markdown)
    """

    @staticmethod
    def read_docx(file_path: str | Path) -> Dict[str, Any]:
        """
        Read .docx file with table support

        Args:
            file_path: Path to .docx file

        Returns:
            Dict with content (markdown), metadata, has_tables flag
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

        logger.info("reading_docx", path=str(file_path))

        try:
            doc = docx.Document(str(file_path))
            content_parts = []
            has_tables = False

            # Iterate through all document elements in order
            for element in doc.element.body:
                # Check if it's a paragraph
                if element.tag.endswith('p'):
                    para = None
                    for p in doc.paragraphs:
                        if p._element == element:
                            para = p
                            break
                    if para and para.text.strip():
                        content_parts.append(para.text.strip())

                # Check if it's a table
                elif element.tag.endswith('tbl'):
                    table = None
                    for t in doc.tables:
                        if t._element == element:
                            table = t
                            break
                    if table:
                        has_tables = True
                        markdown_table = DocumentReader._table_to_markdown(table)
                        content_parts.append(markdown_table)

            content = "\n\n".join(content_parts)

            # Extract metadata from filename
            metadata = DocumentReader._extract_metadata_from_path(file_path)

            logger.info("docx_read_success",
                       path=str(file_path),
                       content_length=len(content),
                       has_tables=has_tables)

            return {
                "content": content,
                "metadata": metadata,
                "has_tables": has_tables,
                "file_path": str(file_path),
                "file_name": file_path.name
            }

        except Exception as e:
            logger.error("docx_read_error", path=str(file_path), error=str(e))
            raise

    @staticmethod
    def _table_to_markdown(table: Table) -> str:
        """
        Convert Word table to Markdown format

        Args:
            table: python-docx Table object

        Returns:
            Markdown formatted table string
        """
        if not table.rows:
            return ""

        markdown_lines = []

        # Extract all rows
        rows_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            rows_data.append(row_data)

        if not rows_data:
            return ""

        # Assume first row is header
        headers = rows_data[0]
        data_rows = rows_data[1:]

        # Build markdown table
        # Header row
        markdown_lines.append("| " + " | ".join(headers) + " |")

        # Separator row
        markdown_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")

        # Data rows
        for row in data_rows:
            # Pad row if necessary to match header length
            while len(row) < len(headers):
                row.append("")
            markdown_lines.append("| " + " | ".join(row[:len(headers)]) + " |")

        return "\n".join(markdown_lines)

    @staticmethod
    def _extract_metadata_from_path(file_path: Path) -> Dict[str, Any]:
        """
        Extract metadata from file path structure

        Example paths:
          base_connaissances/plan_comptable/partie_4/chapitre_7_Comptes combinés.docx
          base_connaissances/actes_uniformes/droit_commercial/titre_3.docx

        Returns:
            Dict with category, section, title, etc.
        """
        parts = file_path.parts
        metadata = {
            "file_path": str(file_path),
            "file_name": file_path.name,
            "file_stem": file_path.stem
        }

        # Find base_connaissances index
        try:
            base_idx = parts.index("base_connaissances")
        except ValueError:
            return metadata

        # Extract category (plan_comptable, actes_uniformes, etc.)
        if len(parts) > base_idx + 1:
            metadata["category"] = parts[base_idx + 1]

        # Extract section (partie_4, droit_commercial, etc.)
        if len(parts) > base_idx + 2:
            metadata["section"] = parts[base_idx + 2]

        # Extract title from filename
        # New simplified format (no duplication):
        #   plan_comptable: chapitre_1 Plan de comptes - subdivisions.docx
        #   actes_uniformes: Livre_1 Statut du commerçant et de l'entreprenant.docx
        #   actes_uniformes: Livre_2_Titre_1 Dispositions générales.docx
        #   actes_uniformes: Preambule ACTE UNIFORME RELATIF AUX CONTRATS.docx

        stem = file_path.stem

        # The filename now contains the full title with proper formatting
        # Just use it as-is (stem without extension)
        metadata["title"] = stem

        return metadata

    @staticmethod
    def read_file(file_path: str | Path) -> Dict[str, Any]:
        """
        Read any supported file format

        Currently supports:
        - .docx (with tables)

        Args:
            file_path: Path to file

        Returns:
            Dict with content and metadata
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()

        if extension == ".docx":
            return DocumentReader.read_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")


# Singleton instance
_document_reader: DocumentReader | None = None


def get_document_reader() -> DocumentReader:
    """Get singleton DocumentReader instance"""
    global _document_reader
    if _document_reader is None:
        _document_reader = DocumentReader()
    return _document_reader

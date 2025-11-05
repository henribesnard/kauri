import os
from pathlib import Path
from docx import Document
import re

def extract_first_title(docx_path):
    """Extrait le premier titre d'un document Word"""
    try:
        doc = Document(docx_path)

        # Chercher le premier paragraphe avec un style de titre
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading') or para.style.name.startswith('Titre'):
                title = para.text.strip()
                if title:
                    return title
            # Si on trouve un paragraphe en gras au début qui pourrait être un titre
            elif para.text.strip() and len(para.runs) > 0:
                if any(run.bold for run in para.runs):
                    title = para.text.strip()
                    if title and len(title) < 200:  # Éviter les longs paragraphes
                        return title

        # Si aucun titre avec style n'est trouvé, chercher le premier texte non vide
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and len(text) < 200:
                return text

        return None
    except Exception as e:
        print(f"Erreur lors de la lecture de {docx_path}: {e}")
        return None

def is_document_empty(docx_path):
    """Vérifie si un document Word est vide"""
    try:
        doc = Document(docx_path)
        for para in doc.paragraphs:
            if para.text.strip():
                return False
        return True
    except Exception as e:
        print(f"Erreur lors de la vérification de {docx_path}: {e}")
        return True

def clean_filename(filename):
    """Nettoie un nom de fichier pour le rendre valide"""
    # Remplacer les caractères interdits
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        filename = filename.replace(char, '')
    # Limiter la longueur
    if len(filename) > 200:
        filename = filename[:200]
    return filename.strip()

def should_process_file(file_path):
    """Vérifie si le fichier doit être traité"""
    path_str = str(file_path)

    # Exclure les dossiers spécifiés
    excluded_dirs = [
        'presentation_ohada',
        'plan_comptable\\partie_3',
        'plan_comptable\\partie_4',
        'plan_comptable/partie_3',
        'plan_comptable/partie_4'
    ]

    for excluded in excluded_dirs:
        if excluded in path_str:
            return False

    return True

def rename_documents(base_path):
    """Renomme tous les documents Word dans base_connaissances"""
    base_path = Path(base_path)

    # Trouver tous les fichiers .docx
    docx_files = list(base_path.rglob('*.docx'))

    renamed_count = 0
    skipped_count = 0
    error_count = 0

    for docx_file in docx_files:
        if not should_process_file(docx_file):
            print(f"Ignoré (dossier exclu): {docx_file}")
            skipped_count += 1
            continue

        # Vérifier si le document est vide
        if is_document_empty(docx_file):
            print(f"Ignoré (vide): {docx_file}")
            skipped_count += 1
            continue

        # Extraire le titre
        title = extract_first_title(docx_file)

        if not title:
            print(f"Aucun titre trouvé: {docx_file}")
            skipped_count += 1
            continue

        # Obtenir le préfixe actuel (nom du fichier sans extension)
        current_name = docx_file.stem

        # Vérifier si le titre est déjà dans le nom du fichier
        title_clean = clean_filename(title)
        if title_clean in current_name or title in current_name:
            print(f"Déjà traité (titre présent): {docx_file}")
            skipped_count += 1
            continue

        # Créer le nouveau nom: préfixe_titre
        new_name = f"{current_name}_{title}"
        new_name = clean_filename(new_name)

        # Créer le nouveau chemin
        new_path = docx_file.parent / f"{new_name}.docx"

        # Vérifier si le fichier existe déjà
        if new_path.exists() and new_path != docx_file:
            print(f"Le fichier existe déjà: {new_path}")
            error_count += 1
            continue

        # Renommer le fichier
        try:
            docx_file.rename(new_path)
            print(f"[OK] Renomme: {docx_file.name} -> {new_path.name}")
            renamed_count += 1
        except Exception as e:
            print(f"[ERREUR] Erreur lors du renommage de {docx_file}: {e}")
            error_count += 1

    print(f"\n=== Résumé ===")
    print(f"Fichiers renommés: {renamed_count}")
    print(f"Fichiers ignorés: {skipped_count}")
    print(f"Erreurs: {error_count}")
    print(f"Total traité: {len(docx_files)}")

if __name__ == "__main__":
    base_connaissances_path = "base_connaissances"
    rename_documents(base_connaissances_path)
